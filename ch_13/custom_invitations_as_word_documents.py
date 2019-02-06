import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

fontName = 'Brush Script Std'
fontSize = 20

doc = docx.Document('guests.docx')
style = doc.styles['Normal']
font = style.font
font.size = Pt(fontSize)
font.name = 'Brush Script Std'

fileStream = open('guests.txt', 'r')


def addInvitation(name):
    paraObj = doc.add_paragraph('')
    paraObj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paraObj.add_run('it would be a pleasure to have the company of\n').\
            italic = True
    paraObj.add_run(name).bold = True
    paraObj.add_run('at 11010 Memor Lane on the Evening\n').italic = True
    paraObj.add_run('April 1st\n')
    paraObj.add_run('at 7 o\'clock\n').italic = True
    doc.add_page_break()
    doc.save('guests.docx')


for line in fileStream.readlines():
    addInvitation(line)
